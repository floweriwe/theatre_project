"""
Сервис модуля документооборота.

Бизнес-логика для работы с документами:
- CRUD операции
- Загрузка и хранение файлов
- Версионирование
- Статистика
- Конвертация DOCX в PDF для предпросмотра
"""
import io
import mimetypes
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import BinaryIO

import magic  # type: ignore
from docx import Document as DocxDocument  # type: ignore
from fastapi import UploadFile
from reportlab.lib.pagesizes import letter  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.exceptions import AlreadyExistsError, NotFoundError, ValidationError
from app.models.document import (
    DocumentCategory,
    Document,
    DocumentVersion,
    Tag,
    DocumentStatus,
    FileType,
)
from app.repositories.document_repository import (
    DocumentCategoryRepository,
    DocumentRepository,
    DocumentVersionRepository,
    TagRepository,
)
from app.schemas.document import (
    DocCategoryCreate,
    DocCategoryUpdate,
    DocumentCreate,
    DocumentUpdate,
    DocumentStats,
    FileUploadResponse,
)


# Маппинг MIME типов на FileType
MIME_TYPE_MAP = {
    "application/pdf": FileType.PDF,
    "application/msword": FileType.DOCUMENT,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCUMENT,
    "application/vnd.oasis.opendocument.text": FileType.DOCUMENT,
    "text/plain": FileType.DOCUMENT,
    "application/vnd.ms-excel": FileType.SPREADSHEET,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": FileType.SPREADSHEET,
    "text/csv": FileType.SPREADSHEET,
    "image/png": FileType.IMAGE,
    "image/jpeg": FileType.IMAGE,
    "image/gif": FileType.IMAGE,
    "image/svg+xml": FileType.IMAGE,
    "image/webp": FileType.IMAGE,
}

# Разрешённые расширения
ALLOWED_EXTENSIONS = {
    ".pdf", ".doc", ".docx", ".odt", ".txt", ".rtf",
    ".xls", ".xlsx", ".csv", ".ods",
    ".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp",
    ".zip", ".rar", ".7z",
}

# Максимальный размер файла (50 MB)
MAX_FILE_SIZE = 50 * 1024 * 1024


class DocumentService:
    """
    Сервис документооборота.
    
    Управляет документами, их версиями, категориями и тегами.
    """
    
    def __init__(self, session: AsyncSession):
        self._session = session
        self._category_repo = DocumentCategoryRepository(session)
        self._document_repo = DocumentRepository(session)
        self._version_repo = DocumentVersionRepository(session)
        self._tag_repo = TagRepository(session)
        self._storage_path = Path(settings.STORAGE_PATH) / "documents"
        self._storage_path.mkdir(parents=True, exist_ok=True)
    
    # =========================================================================
    # File Operations
    # =========================================================================

    def _detect_content_type(self, file_bytes: bytes) -> str:
        """
        Определить MIME-тип файла по его содержимому.

        Использует python-magic для определения типа файла.

        Args:
            file_bytes: Содержимое файла

        Returns:
            MIME-тип (например, 'application/pdf', 'image/jpeg')
        """
        try:
            mime = magic.Magic(mime=True)
            detected_type = mime.from_buffer(file_bytes)
            return detected_type if detected_type else "application/octet-stream"
        except Exception:
            # Fallback на octet-stream если не удалось определить
            return "application/octet-stream"

    def _get_file_type(self, mime_type: str) -> FileType:
        """Определить тип файла по MIME типу."""
        return MIME_TYPE_MAP.get(mime_type, FileType.OTHER)

    def _generate_file_path(self, file_name: str, theater_id: int | None = None) -> str:
        """
        Сгенерировать путь для хранения файла.
        
        Формат: documents/{theater_id}/{year}/{month}/{timestamp}_{filename}
        """
        now = datetime.utcnow()
        theater_dir = str(theater_id) if theater_id else "common"
        
        # Очищаем имя файла
        safe_name = "".join(c for c in file_name if c.isalnum() or c in ".-_")
        unique_name = f"{now.strftime('%Y%m%d_%H%M%S')}_{safe_name}"
        
        relative_path = f"{theater_dir}/{now.year}/{now.month:02d}/{unique_name}"
        return relative_path
    
    async def save_file(
        self,
        file: UploadFile,
        theater_id: int | None = None,
    ) -> FileUploadResponse:
        """
        Сохранить загруженный файл.
        
        Args:
            file: Загружаемый файл
            theater_id: ID театра
            
        Returns:
            Информация о сохранённом файле
            
        Raises:
            ValidationError: Если файл не прошёл валидацию
        """
        # Проверяем расширение
        file_ext = Path(file.filename or "").suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise ValidationError(f"Недопустимый тип файла: {file_ext}")
        
        # Читаем содержимое для проверки размера
        content = await file.read()
        file_size = len(content)
        
        if file_size > MAX_FILE_SIZE:
            raise ValidationError(
                f"Файл слишком большой. Максимум: {MAX_FILE_SIZE // (1024*1024)} MB"
            )
        
        if file_size == 0:
            raise ValidationError("Файл пустой")

        # Определяем MIME тип с помощью python-magic
        mime_type = self._detect_content_type(content)

        # Fallback на заголовок Content-Type или расширение файла
        if mime_type == "application/octet-stream":
            mime_type = file.content_type or mimetypes.guess_type(file.filename or "")[0] or "application/octet-stream"

        file_type = self._get_file_type(mime_type)
        
        # Генерируем путь
        relative_path = self._generate_file_path(file.filename or "file", theater_id)
        full_path = self._storage_path / relative_path
        
        # Создаём директорию
        full_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Сохраняем файл
        with open(full_path, "wb") as f:
            f.write(content)
        
        return FileUploadResponse(
            file_path=relative_path,
            file_name=file.filename or "file",
            file_size=file_size,
            mime_type=mime_type,
            file_type=file_type,
        )
    
    def delete_file(self, file_path: str) -> bool:
        """Удалить файл из хранилища."""
        full_path = self._storage_path / file_path
        if full_path.exists():
            full_path.unlink()
            return True
        return False
    
    def get_file_url(self, file_path: str) -> str:
        """Получить URL для доступа к файлу."""
        return f"/storage/documents/{file_path}"
    
    # =========================================================================
    # Categories
    # =========================================================================
    
    async def get_categories(
        self,
        theater_id: int | None = None,
        skip: int = 0,
        limit: int = 100,
    ) -> tuple[list[DocumentCategory], int]:
        """Получить список категорий."""
        items = await self._category_repo.get_all(skip=skip, limit=limit)
        total = await self._category_repo.count()
        return list(items), total
    
    async def get_categories_tree(
        self,
        theater_id: int | None = None,
    ) -> list[DocumentCategory]:
        """Получить дерево категорий."""
        categories = await self._category_repo.get_tree(theater_id)
        return list(categories)
    
    async def get_category(self, category_id: int) -> DocumentCategory:
        """Получить категорию по ID."""
        category = await self._category_repo.get_by_id(category_id)
        if not category:
            raise NotFoundError(f"Категория с ID {category_id} не найдена")
        return category
    
    async def create_category(
        self,
        data: DocCategoryCreate,
        user_id: int,
        theater_id: int | None = None,
    ) -> DocumentCategory:
        """Создать категорию."""
        # Проверяем уникальность кода
        existing = await self._category_repo.get_by_code(data.code, theater_id)
        if existing:
            raise AlreadyExistsError(f"Категория с кодом '{data.code}' уже существует")
        
        category_data = {
            **data.model_dump(),
            "theater_id": theater_id,
            "created_by_id": user_id,
            "updated_by_id": user_id,
        }

        created = await self._category_repo.create(category_data)
        await self._session.commit()
        
        return await self._category_repo.get_by_id(created.id)
    
    async def update_category(
        self,
        category_id: int,
        data: DocCategoryUpdate,
        user_id: int,
    ) -> DocumentCategory:
        """Обновить категорию."""
        category = await self.get_category(category_id)
        
        update_data = data.model_dump(exclude_unset=True)
        update_data["updated_by_id"] = user_id
        
        # Проверяем код на уникальность
        if "code" in update_data and update_data["code"] != category.code:
            existing = await self._category_repo.get_by_code(update_data["code"], category.theater_id)
            if existing:
                raise AlreadyExistsError(f"Категория с кодом '{update_data['code']}' уже существует")
        
        updated = await self._category_repo.update_by_id(category_id, update_data)
        await self._session.commit()
        
        return updated
    
    async def delete_category(self, category_id: int) -> bool:
        """Удалить категорию (soft delete)."""
        await self.get_category(category_id)
        await self._category_repo.update_by_id(category_id, {"is_active": False})
        await self._session.commit()
        return True
    
    # =========================================================================
    # Tags
    # =========================================================================
    
    async def get_tags(
        self,
        theater_id: int | None = None,
        skip: int = 0,
        limit: int = 100,
    ) -> list[Tag]:
        """Получить список тегов."""
        items = await self._tag_repo.get_all(skip=skip, limit=limit)
        return list(items)
    
    async def create_tag(
        self,
        name: str,
        color: str | None = None,
        theater_id: int | None = None,
    ) -> Tag:
        """Создать тег."""
        existing = await self._tag_repo.get_by_name(name, theater_id)
        if existing:
            raise AlreadyExistsError(f"Тег '{name}' уже существует")
        
        tag = Tag(name=name, color=color, theater_id=theater_id)
        self._session.add(tag)
        await self._session.commit()
        
        return tag
    
    async def delete_tag(self, tag_id: int) -> bool:
        """Удалить тег."""
        tag = await self._tag_repo.get_by_id(tag_id)
        if not tag:
            raise NotFoundError(f"Тег с ID {tag_id} не найден")
        
        await self._session.delete(tag)
        await self._session.commit()
        return True
    
    # =========================================================================
    # Documents
    # =========================================================================
    
    async def get_documents(
        self,
        search: str | None = None,
        category_id: int | None = None,
        status: DocumentStatus | None = None,
        file_type: FileType | None = None,
        tag_ids: list[int] | None = None,
        is_public: bool | None = None,
        is_active: bool | None = True,
        department_id: int | None = None,
        theater_id: int | None = None,
        skip: int = 0,
        limit: int = 20,
    ) -> tuple[list[Document], int]:
        """Получить список документов с фильтрацией."""
        documents, total = await self._document_repo.search(
            search=search,
            category_id=category_id,
            status=status,
            file_type=file_type,
            tag_ids=tag_ids,
            is_public=is_public,
            is_active=is_active,
            department_id=department_id,
            theater_id=theater_id,
            skip=skip,
            limit=limit,
        )
        return list(documents), total
    
    async def get_document(self, document_id: int) -> Document:
        """Получить документ по ID."""
        document = await self._document_repo.get_with_relations(document_id)
        if not document:
            raise NotFoundError(f"Документ с ID {document_id} не найден")
        return document
    
    async def create_document(
        self,
        data: DocumentCreate,
        file_info: FileUploadResponse,
        user_id: int,
        theater_id: int | None = None,
    ) -> Document:
        """
        Создать документ.
        
        Args:
            data: Данные документа
            file_info: Информация о загруженном файле
            user_id: ID пользователя
            theater_id: ID театра
        """
        # Проверяем категорию
        if data.category_id:
            await self.get_category(data.category_id)
        
        # Получаем теги
        tags = []
        if data.tag_ids:
            tags = list(await self._tag_repo.get_by_ids(data.tag_ids))
        
        # Создаём документ
        document = Document(
            name=data.name,
            description=data.description,
            category_id=data.category_id,
            file_path=file_info.file_path,
            file_name=file_info.file_name,
            file_size=file_info.file_size,
            mime_type=file_info.mime_type,
            file_type=file_info.file_type,
            current_version=1,
            status=DocumentStatus.ACTIVE,
            performance_id=data.performance_id,
            is_public=data.is_public,
            theater_id=theater_id,
            created_by_id=user_id,
            updated_by_id=user_id,
            tags=tags,
        )
        
        self._session.add(document)
        await self._session.flush()
        
        # Создаём первую версию
        await self._version_repo.create_version(
            document_id=document.id,
            version=1,
            file_path=file_info.file_path,
            file_name=file_info.file_name,
            file_size=file_info.file_size,
            user_id=user_id,
            comment="Первоначальная загрузка",
        )
        
        await self._session.commit()
        
        return await self._document_repo.get_with_relations(document.id)
    
    async def update_document(
        self,
        document_id: int,
        data: DocumentUpdate,
        user_id: int,
        new_file: FileUploadResponse | None = None,
    ) -> Document:
        """
        Обновить документ.
        
        При загрузке нового файла создаётся новая версия.
        """
        document = await self.get_document(document_id)
        
        update_data = data.model_dump(exclude_unset=True, exclude={"tag_ids"})
        update_data["updated_by_id"] = user_id
        
        # Обновляем теги
        if data.tag_ids is not None:
            tags = list(await self._tag_repo.get_by_ids(data.tag_ids))
            document.tags = tags
        
        # Если есть новый файл — создаём версию
        if new_file:
            new_version = document.current_version + 1
            
            # Сохраняем текущую версию в историю (если ещё не сохранена)
            existing_version = await self._version_repo.get_latest(document_id)
            if not existing_version or existing_version.version != document.current_version:
                await self._version_repo.create_version(
                    document_id=document_id,
                    version=document.current_version,
                    file_path=document.file_path,
                    file_name=document.file_name,
                    file_size=document.file_size,
                    user_id=user_id,
                    comment="Автосохранение перед обновлением",
                )
            
            # Создаём новую версию
            await self._version_repo.create_version(
                document_id=document_id,
                version=new_version,
                file_path=new_file.file_path,
                file_name=new_file.file_name,
                file_size=new_file.file_size,
                user_id=user_id,
                comment="Обновление файла",
            )
            
            # Удаляем старые версии (оставляем 2)
            await self._version_repo.delete_old_versions(document_id, keep_count=2)
            
            # Обновляем данные документа
            update_data["file_path"] = new_file.file_path
            update_data["file_name"] = new_file.file_name
            update_data["file_size"] = new_file.file_size
            update_data["mime_type"] = new_file.mime_type
            update_data["file_type"] = new_file.file_type
            update_data["current_version"] = new_version
        
        await self._document_repo.update_by_id(document_id, update_data)
        await self._session.commit()
        
        return await self._document_repo.get_with_relations(document_id)
    
    async def delete_document(self, document_id: int, user_id: int) -> bool:
        """Удалить документ (soft delete)."""
        await self.get_document(document_id)
        await self._document_repo.update_by_id(document_id, {
            "is_active": False,
            "updated_by_id": user_id,
        })
        await self._session.commit()
        return True
    
    async def archive_document(self, document_id: int, user_id: int) -> Document:
        """Архивировать документ."""
        document = await self.get_document(document_id)
        
        if document.status == DocumentStatus.ARCHIVED:
            raise ValidationError("Документ уже в архиве")
        
        await self._document_repo.update_by_id(document_id, {
            "status": DocumentStatus.ARCHIVED,
            "updated_by_id": user_id,
        })
        await self._session.commit()
        
        return await self._document_repo.get_with_relations(document_id)
    
    async def restore_document(self, document_id: int, user_id: int) -> Document:
        """Восстановить документ из архива."""
        document = await self.get_document(document_id)
        
        if document.status != DocumentStatus.ARCHIVED:
            raise ValidationError("Документ не в архиве")
        
        await self._document_repo.update_by_id(document_id, {
            "status": DocumentStatus.ACTIVE,
            "updated_by_id": user_id,
        })
        await self._session.commit()
        
        return await self._document_repo.get_with_relations(document_id)
    
    # =========================================================================
    # Versions
    # =========================================================================
    
    async def get_document_versions(
        self,
        document_id: int,
        skip: int = 0,
        limit: int = 10,
    ) -> list[DocumentVersion]:
        """Получить версии документа."""
        await self.get_document(document_id)
        versions = await self._version_repo.get_by_document(document_id, skip, limit)
        return list(versions)

    async def get_latest_version(self, document_id: int) -> DocumentVersion:
        """Получить последнюю версию документа."""
        await self.get_document(document_id)
        version = await self._version_repo.get_latest(document_id)
        if not version:
            raise NotFoundError(f"Версии документа {document_id} не найдены")
        return version

    async def get_version_by_id(self, version_id: int) -> DocumentVersion:
        """Получить версию по ID."""
        version = await self._version_repo.get_by_id(version_id)
        if not version:
            raise NotFoundError(f"Версия с ID {version_id} не найдена")
        return version

    async def upload_new_version(
        self,
        document_id: int,
        file: UploadFile,
        user_id: int,
        comment: str | None = None,
    ) -> DocumentVersion:
        """
        Загрузить новую версию документа.

        Создаёт новую версию и обновляет документ.
        Удаляет старые версии, оставляя только 2 последних.
        """
        document = await self.get_document(document_id)

        # Сохраняем файл
        file_info = await self.save_file(file, document.theater_id)

        # Вычисляем новый номер версии
        new_version_num = document.current_version + 1

        # Создаём запись версии
        version = await self._version_repo.create_version(
            document_id=document_id,
            version=new_version_num,
            file_path=file_info.file_path,
            file_name=file_info.file_name,
            file_size=file_info.file_size,
            user_id=user_id,
            comment=comment or "Новая версия",
        )

        # Обновляем документ
        await self._document_repo.update_by_id(document_id, {
            "file_path": file_info.file_path,
            "file_name": file_info.file_name,
            "file_size": file_info.file_size,
            "mime_type": file_info.mime_type,
            "file_type": file_info.file_type,
            "current_version": new_version_num,
            "updated_by_id": user_id,
        })

        # Удаляем старые версии (оставляем 2)
        await self._version_repo.delete_old_versions(document_id, keep_count=2)

        await self._session.commit()

        return version

    # =========================================================================
    # Statistics
    # =========================================================================
    
    async def get_stats(self, theater_id: int | None = None) -> DocumentStats:
        """Получить статистику документов."""
        stats = await self._document_repo.get_stats(theater_id)
        
        # Количество категорий
        categories, _ = await self.get_categories(theater_id)
        stats["categories_count"] = len([c for c in categories if c.is_active])
        
        # Количество тегов
        tags = await self.get_tags(theater_id)
        stats["tags_count"] = len(tags)
        
        return DocumentStats(
            total_documents=stats["total_documents"],
            active=stats.get("active", 0),
            draft=stats.get("draft", 0),
            archived=stats.get("archived", 0),
            total_size=stats["total_size"],
            categories_count=stats["categories_count"],
            tags_count=stats["tags_count"],
            pdf_count=stats.get("pdf_count", 0),
            document_count=stats.get("document_count", 0),
            spreadsheet_count=stats.get("spreadsheet_count", 0),
            image_count=stats.get("image_count", 0),
            other_count=stats.get("other_count", 0),
        )

    # =========================================================================
    # Document Conversion (DOCX to PDF)
    # =========================================================================

    def _get_preview_path(self, document_id: int) -> Path:
        """
        Получить путь для хранения preview PDF.

        Args:
            document_id: ID документа

        Returns:
            Путь к preview PDF файлу
        """
        preview_dir = self._storage_path / "previews"
        preview_dir.mkdir(parents=True, exist_ok=True)
        return preview_dir / f"doc_{document_id}_preview.pdf"

    def _convert_docx_to_pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        """
        Конвертировать DOCX в PDF (простая текстовая версия).

        Извлекает текст из DOCX и создаёт простой PDF с помощью reportlab.
        ОГРАНИЧЕНИЯ: теряется сложное форматирование, таблицы, изображения.

        Args:
            docx_path: Путь к DOCX файлу
            pdf_path: Путь для сохранения PDF

        Returns:
            True если конвертация успешна
        """
        try:
            # Читаем DOCX
            doc = DocxDocument(docx_path)

            # Создаём PDF
            pdf_buffer = io.BytesIO()
            pdf = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch,
            )

            # Стили
            styles = getSampleStyleSheet()
            story = []

            # Заголовок
            title_style = styles["Title"]
            story.append(Paragraph("Document Preview", title_style))
            story.append(Spacer(1, 0.2 * inch))

            # Добавляем параграфы из документа
            normal_style = styles["Normal"]
            for para in doc.paragraphs:
                if para.text.strip():
                    # Очищаем текст от проблемных символов
                    text = para.text.replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(text, normal_style))
                    story.append(Spacer(1, 0.1 * inch))

            # Строим PDF
            pdf.build(story)

            # Сохраняем
            with open(pdf_path, "wb") as f:
                f.write(pdf_buffer.getvalue())

            return True

        except Exception:
            # Если конвертация не удалась, возвращаем False
            return False

    async def get_document_preview_url(self, document_id: int) -> str | None:
        """
        Получить URL для предпросмотра документа в формате PDF.

        Для DOCX/DOC файлов создаёт preview PDF (если ещё не создан).
        Для PDF файлов возвращает URL оригинала.

        Args:
            document_id: ID документа

        Returns:
            URL для предпросмотра или None если конвертация невозможна

        Raises:
            NotFoundError: Если документ не найден
        """
        document = await self.get_document(document_id)

        # Для PDF возвращаем URL оригинала
        if document.mime_type == "application/pdf":
            return self.get_file_url(document.file_path)

        # Проверяем, что это DOCX/DOC
        if document.mime_type not in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]:
            return None

        # Проверяем существование preview
        preview_path = self._get_preview_path(document_id)

        if not preview_path.exists():
            # Создаём preview
            source_path = self._storage_path / document.file_path

            if not source_path.exists():
                raise NotFoundError(f"Исходный файл документа {document_id} не найден")

            # Конвертируем
            success = self._convert_docx_to_pdf(source_path, preview_path)

            if not success:
                return None

        # Возвращаем URL для preview
        relative_path = preview_path.relative_to(self._storage_path)
        return f"/storage/documents/{relative_path}"

    def delete_document_preview(self, document_id: int) -> bool:
        """
        Удалить preview PDF документа.

        Args:
            document_id: ID документа

        Returns:
            True если успешно удалён
        """
        preview_path = self._get_preview_path(document_id)
        if preview_path.exists():
            preview_path.unlink()
            return True
        return False
