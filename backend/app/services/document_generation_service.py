"""
Сервис генерации документов из шаблонов.

Функционал:
- Загрузка шаблона DOCX
- Парсинг плейсхолдеров {{variable}}
- Автозаполнение из связанных сущностей
- Генерация документа с подставленными значениями
- Конвертация в PDF
"""
import io
import re
from datetime import date, datetime
from typing import Any

from docx import Document as DocxDocument  # type: ignore
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.exceptions import NotFoundError, ValidationError
from app.models.document import Document, DocumentStatus, FileType
from app.models.document_template import (
    DocumentTemplate,
    DocumentTemplateVariable,
    VariableType,
)
from app.models.performance import Performance
from app.models.user import User
from app.repositories.template_repository import TemplateRepository
from app.services.minio_service import MinioService


# Паттерн для поиска плейсхолдеров {{variable_name}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*(\w+)\s*\}\}')


class DocumentGenerationService:
    """Сервис генерации документов из шаблонов."""

    def __init__(self, session: AsyncSession) -> None:
        """
        Инициализация сервиса.

        Args:
            session: Асинхронная сессия БД
        """
        self._session = session
        self._template_repo = TemplateRepository(session)
        self._minio = MinioService()

    async def get_template_with_auto_fill(
        self,
        template_id: int,
        performance_id: int | None = None,
    ) -> dict[str, Any]:
        """
        Получить данные шаблона с автозаполненными значениями.

        Args:
            template_id: ID шаблона
            performance_id: ID спектакля для автозаполнения

        Returns:
            Словарь с данными шаблона и автозаполненными значениями
        """
        template = await self._template_repo.get_by_id(template_id)
        if template is None:
            raise NotFoundError(f"Шаблон с ID {template_id} не найден")

        # Загрузка данных для автозаполнения
        performance = None
        if performance_id:
            performance = await self._get_performance(performance_id)

        # Формируем значения переменных
        auto_filled: dict[str, Any] = {}
        suggestions: dict[str, list[dict]] = {}

        for variable in template.variables:
            # Автозаполнение из источника
            if variable.source_field and performance:
                value = self._resolve_source_field(variable.source_field, performance)
                if value is not None:
                    auto_filled[variable.name] = value

            # Подготовка подсказок для разных типов
            if variable.variable_type == VariableType.ACTOR_LIST:
                suggestions[variable.name] = await self._get_actor_suggestions()
            elif variable.variable_type == VariableType.STAFF_LIST:
                suggestions[variable.name] = await self._get_staff_suggestions()
            elif variable.variable_type == VariableType.USER_FIELD:
                suggestions[variable.name] = await self._get_user_suggestions()
            elif variable.variable_type == VariableType.CHOICE and variable.choices:
                suggestions[variable.name] = [
                    {"id": i, "label": choice}
                    for i, choice in enumerate(variable.choices)
                ]

        return {
            "template": template,
            "auto_filled": auto_filled,
            "suggestions": suggestions,
        }

    async def generate_preview(
        self,
        template_id: int,
        variables: dict[str, Any],
        performance_id: int | None = None,
    ) -> bytes:
        """
        Сгенерировать предпросмотр документа.

        Args:
            template_id: ID шаблона
            variables: Значения переменных
            performance_id: ID спектакля

        Returns:
            Байты сгенерированного DOCX
        """
        return await self._generate_document_bytes(
            template_id=template_id,
            variables=variables,
            performance_id=performance_id,
        )

    async def generate_document(
        self,
        template_id: int,
        variables: dict[str, Any],
        performance_id: int | None = None,
        document_name: str | None = None,
        output_format: str = "docx",
        user_id: int | None = None,
        theater_id: int | None = None,
    ) -> Document:
        """
        Сгенерировать документ из шаблона и сохранить.

        Args:
            template_id: ID шаблона
            variables: Значения переменных
            performance_id: ID спектакля
            document_name: Название документа
            output_format: Формат выхода (docx/pdf)
            user_id: ID пользователя
            theater_id: ID театра

        Returns:
            Созданный документ
        """
        template = await self._template_repo.get_by_id(template_id)
        if template is None:
            raise NotFoundError(f"Шаблон с ID {template_id} не найден")

        # Валидация переменных
        await self._validate_variables(template, variables)

        # Генерация DOCX
        docx_bytes = await self._generate_document_bytes(
            template_id=template_id,
            variables=variables,
            performance_id=performance_id,
        )

        # Конвертация в PDF если нужно
        if output_format == "pdf":
            docx_bytes = await self._convert_to_pdf(docx_bytes)
            mime_type = "application/pdf"
            file_ext = ".pdf"
            file_type = FileType.PDF
        else:
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_ext = ".docx"
            file_type = FileType.DOCUMENT

        # Формируем имя файла
        if not document_name:
            document_name = f"{template.name}"
            if performance_id:
                performance = await self._get_performance(performance_id)
                if performance:
                    document_name = f"{template.name} - {performance.title}"

        # Генерация уникального имени файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"{document_name}_{timestamp}{file_ext}"

        # Загрузка в MinIO
        file_path = f"documents/generated/{template.code.lower()}/{file_name}"
        await self._minio.upload_file(
            file_path=file_path,
            content=docx_bytes,
            content_type=mime_type,
        )

        # Создание записи документа
        document = Document(
            name=document_name,
            file_path=file_path,
            file_name=file_name,
            file_size=len(docx_bytes),
            mime_type=mime_type,
            file_type=file_type,
            status=DocumentStatus.ACTIVE,
            performance_id=performance_id,
            generated_from_template_id=template_id,
            generation_data={"variables": variables},
            theater_id=theater_id,
            created_by_id=user_id,
            updated_by_id=user_id,
        )

        self._session.add(document)
        await self._session.flush()
        await self._session.refresh(document)
        await self._session.commit()

        return document

    async def _generate_document_bytes(
        self,
        template_id: int,
        variables: dict[str, Any],
        performance_id: int | None = None,
    ) -> bytes:
        """
        Сгенерировать DOCX документ.

        Args:
            template_id: ID шаблона
            variables: Значения переменных
            performance_id: ID спектакля

        Returns:
            Байты DOCX файла
        """
        template = await self._template_repo.get_by_id(template_id)
        if template is None:
            raise NotFoundError(f"Шаблон с ID {template_id} не найден")

        # Загрузка файла шаблона из MinIO
        template_bytes = await self._minio.download_file(template.file_path)

        # Загрузка данных для автозаполнения недостающих переменных
        performance = None
        if performance_id:
            performance = await self._get_performance(performance_id)

        # Подготовка всех значений
        all_variables = self._prepare_all_variables(
            template=template,
            user_variables=variables,
            performance=performance,
        )

        # Обработка DOCX
        doc = DocxDocument(io.BytesIO(template_bytes))

        # Замена плейсхолдеров в параграфах
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, all_variables)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, all_variables)

        # Замена в headers/footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, all_variables)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, all_variables)

        # Сохранение в байты
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _replace_placeholders_in_paragraph(
        self,
        paragraph: Any,
        variables: dict[str, str],
    ) -> None:
        """
        Заменить плейсхолдеры в параграфе.

        Обрабатывает случаи, когда плейсхолдер разбит на несколько runs.
        """
        # Собираем текст всего параграфа
        full_text = paragraph.text

        # Ищем плейсхолдеры
        matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
        if not matches:
            return

        # Заменяем плейсхолдеры
        new_text = full_text
        for match in reversed(matches):  # Reversed чтобы индексы не сбивались
            var_name = match.group(1)
            if var_name in variables:
                new_text = new_text[:match.start()] + variables[var_name] + new_text[match.end():]

        # Если текст изменился, обновляем параграф
        if new_text != full_text:
            # Сохраняем форматирование первого run
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Очищаем все runs
                for run in paragraph.runs:
                    run.text = ""
                # Устанавливаем новый текст в первый run
                first_run.text = new_text
            else:
                # Если нет runs, просто добавляем текст
                paragraph.clear()
                paragraph.add_run(new_text)

    def _prepare_all_variables(
        self,
        template: DocumentTemplate,
        user_variables: dict[str, Any],
        performance: Performance | None,
    ) -> dict[str, str]:
        """
        Подготовить все значения переменных для подстановки.

        Args:
            template: Шаблон
            user_variables: Значения от пользователя
            performance: Спектакль для автозаполнения

        Returns:
            Словарь переменная -> строковое значение
        """
        result: dict[str, str] = {}

        for variable in template.variables:
            value = None

            # Сначала берём значение от пользователя
            if variable.name in user_variables:
                value = user_variables[variable.name]
            # Затем пробуем автозаполнение
            elif variable.source_field and performance:
                value = self._resolve_source_field(variable.source_field, performance)
            # Затем значение по умолчанию
            elif variable.default_value:
                value = variable.default_value

            # Конвертируем в строку
            if value is not None:
                result[variable.name] = self._format_value(value, variable.variable_type)
            else:
                # Пустая строка для незаполненных
                result[variable.name] = ""

        return result

    def _resolve_source_field(
        self,
        source_field: str,
        performance: Performance,
    ) -> Any:
        """
        Получить значение из источника (performance.field_name).

        Args:
            source_field: Путь к полю, например "performance.title"
            performance: Объект спектакля

        Returns:
            Значение поля или None
        """
        parts = source_field.split(".")
        if len(parts) < 2:
            return None

        entity = parts[0]
        field_path = parts[1:]

        if entity == "performance":
            obj: Any = performance
            for field in field_path:
                if obj is None:
                    return None
                if hasattr(obj, field):
                    obj = getattr(obj, field)
                else:
                    return None
            return obj

        return None

    def _format_value(self, value: Any, variable_type: VariableType) -> str:
        """Форматировать значение в строку."""
        if value is None:
            return ""

        if variable_type == VariableType.DATE:
            if isinstance(value, (date, datetime)):
                return value.strftime("%d.%m.%Y")
            return str(value)

        if variable_type == VariableType.NUMBER:
            if isinstance(value, float):
                return f"{value:,.2f}".replace(",", " ").replace(".", ",")
            return str(value)

        if variable_type in (VariableType.ACTOR_LIST, VariableType.STAFF_LIST):
            if isinstance(value, list):
                return ", ".join(str(v) for v in value)
            return str(value)

        return str(value)

    async def _validate_variables(
        self,
        template: DocumentTemplate,
        variables: dict[str, Any],
    ) -> None:
        """Валидировать переменные."""
        errors = []

        for variable in template.variables:
            if variable.is_required:
                if variable.name not in variables or not variables[variable.name]:
                    # Проверяем, есть ли значение по умолчанию или автозаполнение
                    if not variable.default_value and not variable.source_field:
                        errors.append(f"Поле '{variable.label}' обязательно для заполнения")

        if errors:
            raise ValidationError("; ".join(errors))

    async def _get_performance(self, performance_id: int) -> Performance | None:
        """Получить спектакль по ID."""
        result = await self._session.execute(
            select(Performance).where(Performance.id == performance_id)
        )
        return result.scalar_one_or_none()

    async def _get_actor_suggestions(self) -> list[dict]:
        """Получить список актёров для подсказок."""
        # TODO: Реализовать когда будет модель Actor
        return []

    async def _get_staff_suggestions(self) -> list[dict]:
        """Получить список сотрудников для подсказок."""
        result = await self._session.execute(
            select(User).where(User.is_active == True).limit(100)
        )
        users = result.scalars().all()
        return [
            {"id": u.id, "label": u.full_name or u.email}
            for u in users
        ]

    async def _get_user_suggestions(self) -> list[dict]:
        """Получить список пользователей для подсказок."""
        return await self._get_staff_suggestions()

    async def _convert_to_pdf(self, docx_bytes: bytes) -> bytes:
        """
        Конвертировать DOCX в PDF.

        Uses LibreOffice headless или fallback на reportlab.
        """
        import subprocess
        import tempfile
        import os

        try:
            # Создаём временные файлы
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_file:
                docx_file.write(docx_bytes)
                docx_path = docx_file.name

            # Директория для выхода
            output_dir = tempfile.mkdtemp()

            # Запускаем LibreOffice
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", output_dir,
                    docx_path
                ],
                capture_output=True,
                timeout=60,
            )

            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")

            # Читаем PDF
            pdf_path = os.path.join(
                output_dir,
                os.path.basename(docx_path).replace(".docx", ".pdf")
            )

            with open(pdf_path, "rb") as pdf_file:
                pdf_bytes = pdf_file.read()

            # Очистка
            os.unlink(docx_path)
            os.unlink(pdf_path)
            os.rmdir(output_dir)

            return pdf_bytes

        except (subprocess.TimeoutExpired, FileNotFoundError, RuntimeError) as e:
            # Fallback: возвращаем DOCX как есть с предупреждением
            # В реальной системе здесь должен быть другой способ конвертации
            raise ValidationError(
                f"Конвертация в PDF недоступна. Скачайте документ в формате DOCX. Ошибка: {str(e)}"
            )
